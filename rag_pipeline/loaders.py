"""Document loaders for various file types with structure preservation."""

import json
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional
from datetime import datetime

try:
    from langchain.schema import Document
except ImportError:
    from langchain_core.documents import Document


class BaseDocumentLoader(ABC):
    """Abstract base class for document loaders."""

    def __init__(self, file_path: Path, doc_type: str):
        """Initialize loader."""
        self.file_path = Path(file_path)
        self.doc_type = doc_type
        self.source_name = self.file_path.name

    @abstractmethod
    def load(self) -> list[Document]:
        """Load and parse document into LangChain Document objects."""
        pass

    def _create_document(
        self,
        content: str,
        metadata: Optional[dict] = None,
    ) -> Document:
        """Create a LangChain Document with standard metadata."""
        base_metadata = {
            "source_file": self.source_name,
            "doc_type": self.doc_type,
            "loaded_at": datetime.utcnow().isoformat(),
            "file_path": str(self.file_path),
        }
        if metadata:
            base_metadata.update(metadata)
        return Document(page_content=content, metadata=base_metadata)


class TextDocumentLoader(BaseDocumentLoader):
    """Load plain text files (.txt, .md, .markdown)."""

    def load(self) -> list[Document]:
        """Load text file."""
        content = self.file_path.read_text(encoding="utf-8")
        return [self._create_document(content)]


class PDFDocumentLoader(BaseDocumentLoader):
    """Load PDF files using pdfplumber for better extraction."""

    def load(self) -> list[Document]:
        """Load PDF and extract text per page."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber is required for PDF loading. Install with: pip install pdfplumber")

        documents = []
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    if text.strip():
                        metadata = {
                            "page_number": page_num,
                            "total_pages": len(pdf.pages),
                        }
                        doc = self._create_document(text, metadata)
                        documents.append(doc)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF {self.file_path}: {str(e)}")

        return documents if documents else [self._create_document("")]


class DocxDocumentLoader(BaseDocumentLoader):
    """Load DOCX files (Word documents)."""

    def load(self) -> list[Document]:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        doc = DocxDocument(self.file_path)
        paragraphs = []
        current_section = None

        for element in doc.element.body:
            if element.tag.endswith("p"):
                para = next(
                    (p for p in doc.paragraphs if p._element is element),
                    None,
                )
                if para:
                    style = para.style.name if para.style else "Normal"
                    if "Heading" in style:
                        current_section = para.text
                    if para.text.strip():
                        paragraphs.append(para.text)

        content = "\n\n".join(paragraphs)
        return [self._create_document(content)]


class HTMLDocumentLoader(BaseDocumentLoader):
    """Load HTML files with structure preservation."""

    def load(self) -> list[Document]:
        """Load HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required. Install with: pip install beautifulsoup4")

        html_content = self.file_path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        text_parts = []
        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "li"]):
            text = element.get_text(strip=True)
            if text:
                if element.name.startswith("h"):
                    text_parts.append(f"\n{'#' * int(element.name[1])} {text}\n")
                else:
                    text_parts.append(text)

        content = "\n".join(text_parts)
        return [self._create_document(content)]


class JSONDocumentLoader(BaseDocumentLoader):
    """Load JSON files with hierarchical flattening."""

    def load(self) -> list[Document]:
        """Load JSON and flatten into readable text."""
        try:
            data = json.loads(self.file_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON file {self.file_path}: {str(e)}")

        content = self._flatten_json(data)
        return [self._create_document(content)]

    def _flatten_json(self, obj, prefix=""):
        """Recursively flatten JSON structure into readable text."""
        lines = []

        if isinstance(obj, dict):
            for key, value in obj.items():
                current_path = f"{prefix}.{key}" if prefix else key
                if isinstance(value, (dict, list)):
                    lines.append(f"\n{current_path}:")
                    lines.append(self._flatten_json(value, current_path))
                else:
                    lines.append(f"{current_path}: {value}")

        elif isinstance(obj, list):
            for idx, item in enumerate(obj):
                current_path = f"{prefix}[{idx}]"
                if isinstance(item, (dict, list)):
                    lines.append(f"\n{current_path}:")
                    lines.append(self._flatten_json(item, current_path))
                else:
                    lines.append(f"{current_path}: {item}")

        return "\n".join(filter(None, lines))


class CodeDocumentLoader(BaseDocumentLoader):
    """Load code files with language-aware structure preservation."""

    LANGUAGE_MAP = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".java": "java",
        ".cpp": "cpp",
        ".go": "golang",
        ".rs": "rust",
    }

    def load(self) -> list[Document]:
        """Load code file with language metadata."""
        content = self.file_path.read_text(encoding="utf-8")
        language = self.LANGUAGE_MAP.get(self.file_path.suffix, "unknown")

        metadata = {"language": language}
        return [self._create_document(content, metadata)]


class DocumentLoaderFactory:
    """Factory for creating appropriate document loader based on file type."""

    LOADERS = {
        ".pdf": PDFDocumentLoader,
        ".txt": TextDocumentLoader,
        ".md": TextDocumentLoader,
        ".markdown": TextDocumentLoader,
        ".docx": DocxDocumentLoader,
        ".html": HTMLDocumentLoader,
        ".json": JSONDocumentLoader,
        ".py": CodeDocumentLoader,
        ".js": CodeDocumentLoader,
        ".ts": CodeDocumentLoader,
        ".java": CodeDocumentLoader,
        ".cpp": CodeDocumentLoader,
        ".go": CodeDocumentLoader,
        ".rs": CodeDocumentLoader,
    }

    @classmethod
    def get_loader(cls, file_path: Path) -> BaseDocumentLoader:
        """Get appropriate loader for file type."""
        suffix = file_path.suffix.lower()
        if suffix not in cls.LOADERS:
            raise ValueError(f"Unsupported file type: {suffix}")

        loader_class = cls.LOADERS[suffix]
        return loader_class(file_path, suffix.lstrip("."))

    @classmethod
    def load_document(cls, file_path: Path) -> list[Document]:
        """Load document using appropriate loader."""
        loader = cls.get_loader(file_path)
        return loader.load()
